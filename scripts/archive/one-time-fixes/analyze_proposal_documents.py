#!/usr/bin/env python3
"""
Proposal Document Analyzer
Extracts text from proposal .docx files and analyzes what data could be added to database
WITHOUT making any database changes - generates report for user approval first
"""

import os
import sys
import sqlite3
from pathlib import Path
from docx import Document
import json
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

# Paths
PROPOSAL_FOLDER_1 = "/Users/lukassherman/Library/CloudStorage/OneDrive-Personal/Proposal 2025 (Nung)"
PROPOSAL_FOLDER_2 = "/Users/lukassherman/Library/CloudStorage/OneDrive-Personal/Latest Proposals"
DB_PATH = os.getenv('DATABASE_PATH', './database/bensley_master.db')

def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file"""
    try:
        doc = Document(docx_path)

        # Extract all paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        print(f"Error extracting text from {docx_path}: {e}")
        return None

def get_project_from_database(project_code):
    """Get existing project data from database"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("""
        SELECT *
        FROM proposals
        WHERE project_code = ?
    """, (project_code,))

    project = cursor.fetchone()

    if project:
        project_dict = dict(project)
    else:
        project_dict = None

    conn.close()
    return project_dict

def extract_project_code_from_filename(filename):
    """Extract project code from filename (e.g., '25-024' or '25 BK-033' from filename)"""
    import re
    # Look for pattern like "25-XXX" or "BK-XXX" or "25BK-XXX"
    match = re.search(r'(25-\d{3}|\d{2}\s?BK-\d{3}|\d{2}BK\d{3}|BK-?\d{3})', filename)
    if match:
        code = match.group(1).replace(' ', '')

        # Convert "25-024" to "25BK-024"
        if re.match(r'25-\d{3}', code):
            code = code.replace('25-', '25BK-')
        # Convert "BK024" to "BK-024"
        elif 'BK' in code and '-' not in code:
            code = code.replace('BK', 'BK-')

        # Ensure format is "25BK-033"
        if not code.startswith('25'):
            code = '25' + code
        return code
    return None

def analyze_single_proposal(docx_path, dry_run=True):
    """Analyze a single proposal document"""
    print(f"\n{'='*80}")
    print(f"ANALYZING: {os.path.basename(docx_path)}")
    print(f"{'='*80}")

    # Extract project code from filename
    filename = os.path.basename(docx_path)
    project_code = extract_project_code_from_filename(filename)

    if not project_code:
        print(f"⚠️  Could not extract project code from filename: {filename}")
        return None

    print(f"📋 Project Code: {project_code}")

    # Get existing database entry
    db_project = get_project_from_database(project_code)

    if db_project:
        print(f"✅ Found in database: {db_project['project_name']}")
        print(f"   Status: {db_project['status']}")
        print(f"   Client: {db_project['client_company']}")
        print(f"   Value: {db_project['project_value']}")
    else:
        print(f"⚠️  NOT found in database - this is a NEW project!")

    # Extract text from document
    print(f"\n📄 Extracting text from document...")
    doc_text = extract_text_from_docx(docx_path)

    if not doc_text:
        print(f"❌ Failed to extract text")
        return None

    text_length = len(doc_text)
    print(f"✅ Extracted {text_length:,} characters")

    # Show preview of extracted text
    preview = doc_text[:500].replace('\n', ' ')
    print(f"\n📝 Preview:")
    print(f"   {preview}...")

    return {
        'filename': filename,
        'docx_path': docx_path,
        'project_code': project_code,
        'doc_text': doc_text,
        'doc_text_length': text_length,
        'in_database': db_project is not None,
        'db_project': db_project
    }

def scan_proposal_folders():
    """Scan both proposal folders and return list of .docx files"""
    docx_files = []

    for folder in [PROPOSAL_FOLDER_1, PROPOSAL_FOLDER_2]:
        if not os.path.exists(folder):
            print(f"⚠️  Folder not found: {folder}")
            continue

        print(f"\n📁 Scanning: {folder}")

        for root, dirs, files in os.walk(folder):
            for file in files:
                if file.endswith('.docx') and not file.startswith('~'):  # Skip temp files
                    full_path = os.path.join(root, file)
                    docx_files.append(full_path)

        print(f"   Found {len([f for f in docx_files if folder in f])} .docx files")

    return docx_files

def main():
    print("="*80)
    print("BENSLEY PROPOSAL DOCUMENT ANALYZER")
    print("="*80)
    print("\nThis script will:")
    print("1. Scan proposal folders for .docx files")
    print("2. Extract text from documents")
    print("3. Compare with existing database entries")
    print("4. Generate report of what data COULD be added")
    print("\n⚠️  NO DATABASE CHANGES WILL BE MADE - analysis only!")
    print("="*80)

    # Scan folders
    docx_files = scan_proposal_folders()

    if not docx_files:
        print("\n❌ No .docx files found!")
        return

    print(f"\n✅ Total .docx files found: {len(docx_files)}")

    # Analyze first few as samples
    print(f"\n{'='*80}")
    print("ANALYZING SAMPLE PROPOSALS (first 5)")
    print(f"{'='*80}")

    analyzed = []
    for i, docx_path in enumerate(docx_files[:5]):
        result = analyze_single_proposal(docx_path)
        if result:
            analyzed.append(result)

    # Summary
    print(f"\n{'='*80}")
    print("ANALYSIS SUMMARY")
    print(f"{'='*80}")
    print(f"\nTotal documents analyzed: {len(analyzed)}")
    print(f"Found in database: {sum(1 for a in analyzed if a['in_database'])}")
    print(f"NOT in database (new): {sum(1 for a in analyzed if not a['in_database'])}")

    # Save analysis to JSON
    output_file = "proposal_analysis_results.json"
    with open(output_file, 'w') as f:
        # Convert to serializable format
        serializable = []
        for item in analyzed:
            ser_item = item.copy()
            # Truncate long text for JSON
            if 'doc_text' in ser_item:
                ser_item['doc_text_preview'] = ser_item['doc_text'][:1000]
                del ser_item['doc_text']
            serializable.append(ser_item)

        json.dump(serializable, f, indent=2, default=str)

    print(f"\n✅ Analysis saved to: {output_file}")

    print(f"\n{'='*80}")
    print("NEXT STEPS:")
    print(f"{'='*80}")
    print("1. Review the extracted data above")
    print("2. For projects IN database: we can enrich with more details from proposal docs")
    print("3. For projects NOT in database: we can add as new proposals")
    print("\n⚠️  Waiting for your approval before making any database changes!")

    return analyzed

if __name__ == '__main__':
    results = main()
