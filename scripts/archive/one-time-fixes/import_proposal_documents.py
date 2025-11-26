#!/usr/bin/env python3
"""
Import Proposal Documents
Processes all .docx proposal files and imports/enriches database with:
- Document links
- Scope details
- Historical context for shelved proposals
"""

import os
import sys
import sqlite3
import json
from datetime import datetime
from pathlib import Path
from docx import Document
import openai
from dotenv import load_dotenv
import re

load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')

# Paths
PROPOSAL_FOLDER_1 = "/Users/lukassherman/Library/CloudStorage/OneDrive-Personal/Proposal 2025 (Nung)"
PROPOSAL_FOLDER_2 = "/Users/lukassherman/Library/CloudStorage/OneDrive-Personal/Latest Proposals"
DB_PATH = os.getenv('DATABASE_PATH', './database/bensley_master.db')

def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file"""
    try:
        doc = Document(docx_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        print(f"      ⚠️  Error extracting text: {e}")
        return None

def extract_project_code_from_filename(filename):
    """Extract project code from filename"""
    # Look for pattern like "25-XXX" or "BK-XXX" or "25BK-XXX"
    match = re.search(r'(25-\d{3}|\d{2}\s?BK-\d{3}|\d{2}BK\d{3}|BK-?\d{3})', filename)
    if match:
        code = match.group(1).replace(' ', '')

        # Convert "25-024" to "25BK-024"
        if re.match(r'25-\d{3}', code):
            code = code.replace('25-', '25BK-')
        # Convert "BK024" to "BK-024"
        elif 'BK' in code and '-' not in code:
            code = code.replace('BK', 'BK-')

        # Ensure format is "25BK-033" or "BK-033"
        if not code.startswith('25') and not code.startswith('BK'):
            code = '25' + code

        return code
    return None

def extract_data_with_ai(doc_text, project_code):
    """Use OpenAI to extract structured data from proposal document"""

    # Truncate text if too long (max 15k chars for API)
    doc_text_truncated = doc_text[:15000] if len(doc_text) > 15000 else doc_text

    prompt = f"""Extract key data from this Bensley Design Studios proposal document.

Return JSON with these fields (use null if not found):
{{
  "project_name": "Full project name",
  "client_company": "Client company name",
  "location": "City, Country",
  "contract_value": numeric value only (no currency symbols),
  "currency": "USD, THB, etc.",
  "proposal_date": "YYYY-MM-DD",
  "services": ["landscape", "architecture", "interior"] (which services included),
  "scope_summary": "Brief 1-2 sentence summary of scope",
  "payment_terms": "Brief summary of payment terms"
}}

Document text:
{doc_text_truncated}

Project Code: {project_code}
"""

    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",  # Cheaper model for batch processing
            messages=[
                {"role": "system", "content": "Extract data accurately and return valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=800
        )

        extracted_text = response.choices[0].message.content

        # Try to extract JSON from response
        if "```json" in extracted_text:
            json_start = extracted_text.find("```json") + 7
            json_end = extracted_text.find("```", json_start)
            json_str = extracted_text[json_start:json_end].strip()
        elif "```" in extracted_text:
            json_start = extracted_text.find("```") + 3
            json_end = extracted_text.find("```", json_start)
            json_str = extracted_text[json_start:json_end].strip()
        else:
            json_str = extracted_text

        return json.loads(json_str)

    except Exception as e:
        print(f"      ⚠️  AI extraction error: {e}")
        return None

def check_project_exists(cursor, project_code):
    """Check if project exists in database"""
    # Try exact match first
    cursor.execute("""
        SELECT proposal_id, project_name, status, document_path
        FROM proposals
        WHERE project_code = ?
    """, (project_code,))

    row = cursor.fetchone()
    if row:
        return {
            'proposal_id': row[0],
            'project_name': row[1],
            'status': row[2],
            'document_path': row[3]
        }

    # Try without '25' prefix (e.g., "BK-017" instead of "25BK-017")
    if project_code.startswith('25BK-'):
        alt_code = project_code.replace('25BK-', 'BK-')
        cursor.execute("""
            SELECT proposal_id, project_name, status, document_path
            FROM proposals
            WHERE project_code = ?
        """, (alt_code,))

        row = cursor.fetchone()
        if row:
            return {
                'proposal_id': row[0],
                'project_name': row[1],
                'status': row[2],
                'document_path': row[3]
            }

    return None

def import_or_update_proposal(cursor, project_code, docx_path, extracted_data, dry_run=True):
    """Import new proposal or update existing one"""

    # Check if exists
    existing = check_project_exists(cursor, project_code)

    if existing:
        # Update existing project with document link and enriched data
        print(f"      ✅ Found in DB: {existing['project_name']}")

        if dry_run:
            print(f"      [DRY RUN] Would update with document_path and enriched data")
            return 'updated_dry'

        # Build update query for fields that have values
        updates = []
        params = []

        if extracted_data:
            if not existing['document_path']:
                updates.append("document_path = ?")
                params.append(docx_path)

            if extracted_data.get('location'):
                updates.append("location = ?")
                params.append(extracted_data['location'])

            if extracted_data.get('currency'):
                updates.append("currency = ?")
                params.append(extracted_data['currency'])

            if extracted_data.get('proposal_date'):
                updates.append("proposal_sent_date = ?")
                params.append(extracted_data['proposal_date'])

            if extracted_data.get('scope_summary'):
                updates.append("scope_summary = ?")
                params.append(extracted_data['scope_summary'])

            if extracted_data.get('payment_terms'):
                updates.append("payment_terms = ?")
                params.append(extracted_data['payment_terms'])

            # Update service flags
            if extracted_data.get('services'):
                services = extracted_data['services']
                if 'landscape' in services:
                    updates.append("is_landscape = 1")
                if 'architecture' in services or 'architectural' in services:
                    updates.append("is_architect = 1")
                if 'interior' in services:
                    updates.append("is_interior = 1")

        if updates:
            updates.append("updated_at = datetime('now')")
            params.append(existing['proposal_id'])

            query = f"UPDATE proposals SET {', '.join(updates)} WHERE proposal_id = ?"
            cursor.execute(query, params)

        return 'updated'

    else:
        # Insert new proposal as 'shelved' status
        print(f"      🆕 New proposal - adding as 'shelved'")

        if dry_run:
            print(f"      [DRY RUN] Would insert new proposal")
            return 'inserted_dry'

        if not extracted_data:
            print(f"      ⚠️  No extracted data - skipping")
            return 'skipped'

        # Prepare insert data
        project_name = extracted_data.get('project_name') or f"Project {project_code}"
        client_company = extracted_data.get('client_company')
        location = extracted_data.get('location')
        project_value = extracted_data.get('contract_value')
        currency = extracted_data.get('currency', 'USD')
        proposal_sent_date = extracted_data.get('proposal_date')
        scope_summary = extracted_data.get('scope_summary')
        payment_terms = extracted_data.get('payment_terms')

        services = extracted_data.get('services', [])
        is_landscape = 1 if 'landscape' in services else 0
        is_architect = 1 if 'architecture' in services or 'architectural' in services else 0
        is_interior = 1 if 'interior' in services else 0

        cursor.execute("""
            INSERT INTO proposals
            (project_code, project_name, client_company, location, project_value,
             currency, proposal_sent_date, scope_summary, payment_terms,
             is_landscape, is_architect, is_interior, status, document_path,
             created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'shelved', ?, datetime('now'), datetime('now'))
        """, (project_code, project_name, client_company, location, project_value,
              currency, proposal_sent_date, scope_summary, payment_terms,
              is_landscape, is_architect, is_interior, docx_path))

        return 'inserted'

def main():
    print("="*80)
    print("BENSLEY PROPOSAL DOCUMENTS IMPORTER")
    print("="*80)

    # Check command line args
    import sys
    dry_run = '--dry-run' in sys.argv or len(sys.argv) == 1  # Default to dry-run if no args

    # Scan for all .docx files
    docx_files = []
    for folder in [PROPOSAL_FOLDER_1, PROPOSAL_FOLDER_2]:
        if os.path.exists(folder):
            for root, dirs, files in os.walk(folder):
                for file in files:
                    if file.endswith('.docx') and not file.startswith('~'):
                        docx_files.append(os.path.join(root, file))

    print(f"\n📄 Found {len(docx_files)} proposal documents")

    # Show what will happen
    print(f"\nThis will:")
    print(f"  1. Extract data from all {len(docx_files)} documents using AI")
    print(f"  2. Update existing proposals with document links and enriched data")
    print(f"  3. Add new proposals as 'shelved' status for historical context")
    print(f"  4. Cost: ~${len(docx_files) * 0.02:.2f} (using GPT-4o-mini)")
    print(f"\nMode: {'DRY RUN' if dry_run else 'LIVE IMPORT'}")
    print(f"(Use --live flag to run live import)")

    # Connect to database
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    # Process documents
    print(f"\n{'='*80}")
    print(f"PROCESSING DOCUMENTS {'[DRY RUN]' if dry_run else '[LIVE]'}")
    print(f"{'='*80}\n")

    stats = {
        'updated': 0,
        'inserted': 0,
        'skipped': 0,
        'errors': 0
    }

    for i, docx_path in enumerate(docx_files, 1):
        filename = os.path.basename(docx_path)
        print(f"[{i}/{len(docx_files)}] {filename}")

        # Extract project code
        project_code = extract_project_code_from_filename(filename)
        if not project_code:
            print(f"      ⚠️  Could not extract project code - skipping")
            stats['skipped'] += 1
            continue

        print(f"      📋 Project: {project_code}")

        # Extract text
        doc_text = extract_text_from_docx(docx_path)
        if not doc_text or len(doc_text) < 500:
            print(f"      ⚠️  Insufficient text extracted - skipping")
            stats['skipped'] += 1
            continue

        # AI extraction
        print(f"      🤖 Extracting data...")
        extracted_data = extract_data_with_ai(doc_text, project_code)

        if not extracted_data:
            print(f"      ⚠️  AI extraction failed - skipping")
            stats['errors'] += 1
            continue

        # Import or update
        result = import_or_update_proposal(cursor, project_code, docx_path, extracted_data, dry_run)

        if result in ['updated', 'updated_dry']:
            stats['updated'] += 1
        elif result in ['inserted', 'inserted_dry']:
            stats['inserted'] += 1
        else:
            stats['skipped'] += 1

        print()

    # Commit if not dry run
    if not dry_run:
        conn.commit()
        print(f"\n✅ Changes committed to database")
    else:
        print(f"\n[DRY RUN] No changes made to database")

    conn.close()

    # Summary
    print(f"\n{'='*80}")
    print("IMPORT SUMMARY")
    print(f"{'='*80}")
    print(f"Total documents processed: {len(docx_files)}")
    print(f"  Updated existing: {stats['updated']}")
    print(f"  Inserted new (shelved): {stats['inserted']}")
    print(f"  Skipped: {stats['skipped']}")
    print(f"  Errors: {stats['errors']}")

if __name__ == '__main__':
    main()
