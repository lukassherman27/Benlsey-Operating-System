#!/usr/bin/env python3
"""
Intelligent Contract Parser
Extracts contract details from .docx and .pdf files and stores in database

Flexibly handles:
- Multiple disciplines (Landscape, Architecture, Interior, Branding, etc.)
- Multiple phases (Mobilization, Schematic, Conceptual, DD, CD, CA, etc.)
- Variable contract structures
"""

import sqlite3
import re
from pathlib import Path
from datetime import datetime
from docx import Document
import anthropic
import os
import json

DB_PATH = "database/bensley_master.db"

# Standard phase order mapping
PHASE_ORDER = {
    'mobilization': 1,
    'schematic': 2,
    'conceptual': 3,
    'design development': 4,
    'construction documents': 5,
    'construction observation': 6,
    'construction administration': 6,  # Same as observation
}

class ContractParser:
    def __init__(self, contract_file_path):
        self.contract_file = Path(contract_file_path)
        self.raw_text = ""
        self.parsed_data = {}

        # Initialize Anthropic client
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if api_key:
            self.client = anthropic.Anthropic(api_key=api_key)
        else:
            self.client = None
            print("⚠️  No ANTHROPIC_API_KEY found - will use rule-based parsing only")

    def extract_text_from_docx(self):
        """Extract text from .docx file"""
        print(f"📄 Reading {self.contract_file.name}...")

        doc = Document(self.contract_file)
        text_lines = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_lines.append(" | ".join(row_text))

        self.raw_text = "\n".join(text_lines)
        return self.raw_text

    def parse_with_ai(self):
        """Use Claude AI to intelligently parse contract"""
        if not self.client:
            print("⚠️  Skipping AI parsing (no API key)")
            return None

        print("🤖 Parsing contract with AI...")

        prompt = f"""Analyze this Bensley Design Group contract and extract the following information in JSON format:

{{
  "project_code": "XX BK-XXX",
  "client_name": "Client company name",
  "contract_date": "YYYY-MM-DD",
  "contract_term_months": 24,
  "total_contract_value": 3450000.00,
  "payment_due_days": 30,
  "late_payment_interest_rate": 1.5,

  "disciplines": [
    {{
      "name": "Landscape Architecture",
      "total_fee": 862500.00,
      "phases": [
        {{
          "name": "Mobilization Fee",
          "fee": 129375.00,
          "order": 1
        }},
        {{
          "name": "Conceptual Design",
          "fee": 215625.00,
          "order": 3
        }}
      ]
    }}
  ]
}}

Extract ALL disciplines mentioned (Landscape, Architecture, Interior, Branding, etc.) and ALL their phases.
Common phases: Mobilization, Schematic Design (if mentioned), Conceptual Design, Design Development, Construction Documents, Construction Observation/Administration.

Contract text:
{self.raw_text[:15000]}

Return ONLY valid JSON, no explanation."""

        try:
            message = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )

            response_text = message.content[0].text.strip()

            # Extract JSON from response (handle markdown code blocks)
            if "```json" in response_text:
                json_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_text = response_text.split("```")[1].split("```")[0].strip()
            else:
                json_text = response_text

            parsed = json.loads(json_text)
            self.parsed_data = parsed

            print(f"✅ Parsed contract: {parsed.get('project_code', 'Unknown')}")
            print(f"   Total value: ${parsed.get('total_contract_value', 0):,.2f}")
            print(f"   Disciplines: {len(parsed.get('disciplines', []))}")

            return parsed

        except Exception as e:
            print(f"❌ AI parsing failed: {e}")
            return None

    def parse_with_rules(self):
        """Fallback rule-based parsing"""
        print("📋 Using rule-based parsing...")

        parsed = {
            'project_code': None,
            'total_contract_value': 0,
            'disciplines': []
        }

        # Extract project code
        code_match = re.search(r'(\d{2}\s+BK-\d{3})', self.raw_text)
        if code_match:
            parsed['project_code'] = code_match.group(1)

        # Extract fees
        fee_pattern = r'(USD|usd|\$)\s*[\s]*([0-9,]+(?:\.[0-9]{2})?)'
        fees = re.findall(fee_pattern, self.raw_text)

        if fees:
            # Find largest fee as total contract value
            amounts = [float(f[1].replace(',', '')) for f in fees]
            parsed['total_contract_value'] = max(amounts)

        return parsed

    def save_to_database(self):
        """Save parsed contract data to database"""
        if not self.parsed_data:
            print("⚠️  No data to save")
            return False

        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()

        try:
            # Find project by code
            project_code = self.parsed_data.get('project_code')
            if not project_code:
                print("❌ No project code found")
                return False

            cursor.execute("SELECT project_id FROM projects WHERE project_code = ?", (project_code,))
            result = cursor.fetchone()

            if not result:
                print(f"⚠️  Project {project_code} not found in database")
                return False

            project_id = result[0]
            print(f"✅ Found project ID: {project_id}")

            # Save contract metadata with provenance
            cursor.execute("""
                INSERT OR REPLACE INTO contract_metadata (
                    project_id,
                    contract_date,
                    contract_term_months,
                    contract_file_path,
                    total_contract_value_usd,
                    total_landscape_fee_usd,
                    total_architecture_fee_usd,
                    total_interior_fee_usd,
                    total_branding_fee_usd,
                    payment_due_days,
                    late_payment_interest_rate,
                    client_name,
                    parsed_at,
                    parsed_by,
                    source_type,
                    source_reference,
                    created_by
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                project_id,
                self.parsed_data.get('contract_date'),
                self.parsed_data.get('contract_term_months'),
                str(self.contract_file),
                self.parsed_data.get('total_contract_value'),
                self._get_discipline_total('Landscape'),
                self._get_discipline_total('Architecture') or self._get_discipline_total('Architectural'),
                self._get_discipline_total('Interior'),
                self._get_discipline_total('Branding'),
                self.parsed_data.get('payment_due_days', 30),
                self.parsed_data.get('late_payment_interest_rate'),
                self.parsed_data.get('client_name'),
                datetime.now().isoformat(),
                'ai',
                'ai',
                f'contract_file:{self.contract_file.name}',
                'parse_contracts_ai'
            ))

            # Update project total fee
            cursor.execute("""
                UPDATE projects
                SET total_fee_usd = ?
                WHERE project_id = ?
            """, (self.parsed_data.get('total_contract_value'), project_id))

            # Delete existing phases for this project (to avoid duplicates on re-parse)
            cursor.execute("DELETE FROM contract_phases WHERE project_id = ?", (project_id,))

            # Save contract phases
            phases_saved = 0
            for discipline in self.parsed_data.get('disciplines', []):
                disc_name = discipline['name']

                for phase in discipline.get('phases', []):
                    phase_name = phase['name']
                    phase_fee = phase.get('fee', 0)
                    phase_order = phase.get('order', 0)

                    # Normalize phase order if not provided
                    if not phase_order:
                        phase_name_lower = phase_name.lower()
                        for key, order in PHASE_ORDER.items():
                            if key in phase_name_lower:
                                phase_order = order
                                break

                    cursor.execute("""
                        INSERT INTO contract_phases (
                            project_id,
                            discipline,
                            phase_name,
                            phase_order,
                            phase_fee_usd,
                            status
                        ) VALUES (?, ?, ?, ?, ?, ?)
                    """, (
                        project_id,
                        disc_name,
                        phase_name,
                        phase_order,
                        phase_fee,
                        'pending'
                    ))
                    phases_saved += 1

            conn.commit()
            print(f"✅ Saved {phases_saved} contract phases")
            return True

        except Exception as e:
            print(f"❌ Database error: {e}")
            conn.rollback()
            return False
        finally:
            conn.close()

    def _get_discipline_total(self, discipline_keyword):
        """Get total fee for a discipline by keyword match"""
        for disc in self.parsed_data.get('disciplines', []):
            if discipline_keyword.lower() in disc['name'].lower():
                return disc.get('total_fee')
        return None

    def process(self):
        """Main processing pipeline"""
        print("\n" + "=" * 100)
        print(f"PROCESSING CONTRACT: {self.contract_file.name}")
        print("=" * 100)

        # Step 1: Extract text
        if self.contract_file.suffix.lower() == '.docx':
            self.extract_text_from_docx()
        else:
            print(f"❌ Unsupported file type: {self.contract_file.suffix}")
            return False

        # Step 2: Parse with AI
        self.parse_with_ai()

        # Step 3: Fallback to rules if AI failed
        if not self.parsed_data:
            self.parsed_data = self.parse_with_rules()

        # Step 4: Save to database
        if self.parsed_data:
            return self.save_to_database()

        return False


def parse_contract(contract_file_path):
    """Parse a single contract file"""
    parser = ContractParser(contract_file_path)
    return parser.process()


def parse_all_contracts_in_folder(folder_path):
    """Parse all contract files in a folder"""
    folder = Path(folder_path)

    if not folder.exists():
        print(f"❌ Folder not found: {folder_path}")
        return

    # Find all .docx files
    contract_files = list(folder.glob("*.docx")) + list(folder.glob("*.pdf"))

    if not contract_files:
        print(f"⚠️  No contract files found in {folder_path}")
        return

    print("\n" + "=" * 100)
    print(f"FOUND {len(contract_files)} CONTRACT FILES")
    print("=" * 100)

    success_count = 0
    for contract_file in contract_files:
        try:
            if parse_contract(contract_file):
                success_count += 1
        except Exception as e:
            print(f"❌ Error processing {contract_file.name}: {e}")

    print("\n" + "=" * 100)
    print(f"COMPLETED: {success_count}/{len(contract_files)} contracts parsed successfully")
    print("=" * 100)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python3 parse_contracts.py <contract_file_or_folder>")
        print("\nExamples:")
        print("  python3 parse_contracts.py /Users/lukassherman/Desktop/contract.docx")
        print("  python3 parse_contracts.py /Users/lukassherman/Desktop/Contracts/")
        sys.exit(1)

    path = sys.argv[1]
    path_obj = Path(path)

    if path_obj.is_file():
        parse_contract(path)
    elif path_obj.is_dir():
        parse_all_contracts_in_folder(path)
    else:
        print(f"❌ Invalid path: {path}")
