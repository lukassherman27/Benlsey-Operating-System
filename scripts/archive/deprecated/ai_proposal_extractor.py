#!/usr/bin/env python3
"""
AI-Powered Proposal Data Extractor
Uses OpenAI to extract structured data from proposal documents
"""

import os
import sys
import json
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
import openai

load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')

def extract_text_from_docx(docx_path):
    """Extract all text from a .docx file"""
    doc = Document(docx_path)

    # Extract all paragraphs
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\n".join(full_text)

def extract_data_with_ai(doc_text, project_code):
    """Use OpenAI to extract structured data from proposal document"""

    prompt = f"""You are analyzing a Bensley Design Studios contract/proposal document.

Extract ALL relevant structured data from this document. Focus on:

1. PROJECT IDENTIFICATION:
   - Project code (e.g., 25BK-024)
   - Project name
   - Client company name
   - Client contact person
   - Client address
   - Location (city, country)

2. CONTRACT DETAILS:
   - Contract/agreement date (when it was signed)
   - Contract value/fee (total USD amount)
   - Currency
   - Payment terms/milestones
   - Payment schedule

3. SCOPE OF WORK:
   - Services included (landscape, architecture, interior)
   - Project phases
   - Deliverables
   - Timeline/duration

4. PROJECT INFORMATION:
   - Project type (hotel, resort, residential, etc.)
   - Site area
   - Number of rooms/units (if applicable)
   - Budget range (if mentioned)

5. KEY DATES:
   - Design phase dates
   - Expected completion dates
   - Milestone dates

6. STATUS INDICATORS:
   - Is this a signed contract or just a proposal?
   - Any terms about revisions/changes?

Return the data as a JSON object with clear keys. Use null for missing values.
For dates, use ISO format (YYYY-MM-DD).
For amounts, extract numeric value only.

Document text:
{doc_text[:15000]}
... (truncated for token limit)

Project Code: {project_code}
"""

    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a data extraction specialist for architecture/design contracts. Extract structured data accurately."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,  # Low temperature for consistent extraction
            max_tokens=2000
        )

        # Parse the JSON response
        extracted_text = response.choices[0].message.content

        # Try to extract JSON from response (might have markdown formatting)
        if "```json" in extracted_text:
            json_start = extracted_text.find("```json") + 7
            json_end = extracted_text.find("```", json_start)
            json_str = extracted_text[json_start:json_end].strip()
        elif "```" in extracted_text:
            json_start = extracted_text.find("```") + 3
            json_end = extracted_text.find("```", json_start)
            json_str = extracted_text[json_start:json_end].strip()
        else:
            json_str = extracted_text

        extracted_data = json.loads(json_str)
        return extracted_data

    except Exception as e:
        print(f"Error during AI extraction: {e}")
        return None

def analyze_document(docx_path, project_code):
    """Full analysis of a single document"""
    print(f"\n{'='*80}")
    print(f"AI EXTRACTION: {os.path.basename(docx_path)}")
    print(f"{'='*80}")
    print(f"Project Code: {project_code}")

    # Extract text
    print(f"\n📄 Extracting text...")
    doc_text = extract_text_from_docx(docx_path)
    print(f"✅ Extracted {len(doc_text):,} characters")

    # AI extraction
    print(f"\n🤖 Running AI extraction (OpenAI GPT-4o)...")
    extracted_data = extract_data_with_ai(doc_text, project_code)

    if not extracted_data:
        print(f"❌ AI extraction failed")
        return None

    print(f"✅ AI extraction complete!")

    # Display results
    print(f"\n{'='*80}")
    print("EXTRACTED DATA:")
    print(f"{'='*80}")
    print(json.dumps(extracted_data, indent=2))

    return {
        'project_code': project_code,
        'filename': os.path.basename(docx_path),
        'docx_path': docx_path,
        'extracted_data': extracted_data,
        'extraction_timestamp': datetime.now().isoformat()
    }

def main():
    """Test AI extraction on sample document"""
    print("="*80)
    print("AI-POWERED PROPOSAL DATA EXTRACTOR")
    print("="*80)

    # Test with the Four Seasons Chiang Rai document
    test_file = "/Users/lukassherman/Library/CloudStorage/OneDrive-Personal/Proposal 2025 (Nung)/25-024 Four Seasons Tented Camp Chiangrai, Thailand-Design Maintenance (2025) revised 17 Apr 25.docx"
    project_code = "25BK-024"

    if not os.path.exists(test_file):
        print(f"❌ Test file not found: {test_file}")
        return

    result = analyze_document(test_file, project_code)

    if result:
        # Save results
        output_file = "ai_extraction_sample.json"
        with open(output_file, 'w') as f:
            json.dump(result, f, indent=2)
        print(f"\n✅ Results saved to: {output_file}")

        print(f"\n{'='*80}")
        print("DATABASE MAPPING SUGGESTIONS:")
        print(f"{'='*80}")

        extracted = result['extracted_data']

        print("\nThese fields could be added to 'proposals' table:")
        print(f"  - project_code: {extracted.get('project_code')}")
        print(f"  - project_name: {extracted.get('project_name')}")
        print(f"  - client_company: {extracted.get('client_company')}")
        print(f"  - client_contact: {extracted.get('client_contact_person')}")
        print(f"  - project_value: {extracted.get('contract_value_usd')}")
        print(f"  - location: {extracted.get('location')}")
        print(f"  - services: {extracted.get('services_included')}")
        print(f"  - signed_date: {extracted.get('contract_date')}")
        print(f"  - status: 'active' (since this is a signed contract)")

        print(f"\n⚠️  This is ONE document out of 139 total documents!")
        print(f"⚠️  All extracted data shown above is ready to import.")
        print(f"\n🚦 Waiting for your approval before any database changes!")

if __name__ == '__main__':
    main()
